from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches, Cm
from indexer.search import find_osnovni_oblik
from .models import *
from .utils import *


def all_words_from_all_pubs_docx(filepath):
    document = Document()
    styles = document.styles
    style = styles.add_style('Stavka', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Normal']
    font = style.font
    font.name = 'Dijakritika'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.left_indent = Cm(1)
    paragraph_format.first_line_indent = Cm(-1)
    paragraph_format.space_after = Pt(6)
    paragraph_format.widow_control = False
    words = all_words_from_all_pubs()
    merge_all_forms(words)
    for word, entry in words.items():
        paragraph = document.add_paragraph()
        paragraph.style = style
        rec_run = paragraph.add_run(word)
        rec_run.bold = True
        paragraph.add_run(f': {entry["freq"]}')
        for pub in entry['pubs']:
            paragraph.add_run('\n')
            pubskr = paragraph.add_run(pub['pubskr'])
            pubskr.italic = True
            paragraph.add_run(f': {pub["freq"]}')
    document.save(filepath)


def merge_all_forms(words):
    for_deletion = []
    all_words = list(words.keys())
    for word in all_words:
        entry = words[word]
        rec = entry['word']
        osnovni_oblici = find_osnovni_oblik(rec)
        if len(osnovni_oblici) != 1:
            continue
        oo = osnovni_oblici[0]['rec']
        if oo == rec:
            continue
        try:
            osnob = words[oo]
        except KeyError:
            osnob = {'word': oo, 'pubs': [], 'freq': 0}
            words[oo] = osnob
        osnob['pubs'].extend(entry['pubs'])
        osnob['freq'] += entry['freq']
        for_deletion.append(rec)
    for dd in for_deletion:
        del words[dd]


def all_words_from_all_pubs():
    words = {}
    for pub in Publikacija.objects.all():
        pub_words = all_words_from_pub(pub)
        for word, item in pub_words.items():
            if words.get(word):
                words[word]['pubs'].append(item)
            else:
                words[word] = {'word': word, 'pubs': [item], 'freq': 0}
    for word, entry in words.items():
        entry['freq'] = sum(item['freq'] for item in entry['pubs'])   
    return words


def all_words_from_pub(pub):
    words = {}
    try:
        tekst = TekstPublikacije.objects.filter(publikacija=pub).order_by('redni_broj')
        for t in tekst:
            page_text = t.tekst
            parsed_words = parse_to_words(page_text)
            for w in parsed_words:
                w = w.lower()
                entry = {
                    'word': w,
                    'pubid': pub.id,
                    'pubskr': pub.skracenica,
                    'pages': [t.redni_broj],
                    'freq': 1
                }
                if words.get(w):
                    words[w]['freq'] += 1
                    if t.redni_broj not in words[w]['pages']:
                        words[w]['pages'].append(t.redni_broj)
                else:
                    words[w] = entry
        return words
    except TekstPublikacije.DoesNotExist:
        return {}


def parse_to_words(text):
    text = remove_punctuation(text)
    words = text.split()
    return [w for w in words if len(w) > 0]
